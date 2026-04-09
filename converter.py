"""
Markdown → DOCX converter.

Uses mistune ≥ 3.0 for parsing and python-docx for document generation.
All markdown elements are preserved; no content is silently dropped.
"""

import io
import re
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import mistune
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

# ─── XML / colour helpers ─────────────────────────────────────────────────────

def _hex(color: str) -> RGBColor:
    h = color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_font(run, font_name: str) -> None:
    """Set font name for both Latin (ascii/hAnsi) and East Asian (eastAsia) ranges.

    python-docx's run.font.name only sets w:ascii and w:hAnsi.
    Without w:eastAsia, Word falls back to the document's default CJK font
    (often SimSun), which causes Korean syllables to decompose into Jamo.
    """
    run.font.name = font_name  # sets w:ascii + w:hAnsi
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def _clean_hex(color: str) -> str:
    return color.lstrip("#").upper()


def _cell_shade(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _clean_hex(fill))
    tcPr.append(shd)


def _para_shade(para, fill: str):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _clean_hex(fill))
    pPr.append(shd)


def _outline_level(para, level: int):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(old)
    el = OxmlElement("w:outlineLvl")
    el.set(qn("w:val"), str(max(0, level - 1)))
    pPr.append(el)


def _add_hyperlink(para, text: str, url: str):
    """Embed a proper hyperlink relationship in the paragraph."""
    try:
        r_id = para.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), r_id)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        hl.append(r)
        para._p.append(hl)
    except Exception:
        run = para.add_run(f"{text} ({url})")
        run.font.underline = True
        run.font.color.rgb = RGBColor(0x00, 0x56, 0xD6)


def _para_left_border(para, color: str = "2E75B6", sz: str = "24"):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), _clean_hex(color))
    pBdr.append(left)
    pPr.append(pBdr)


def _remove_numbering(para) -> None:
    """Remove any automatic bullet/numbering (w:numPr) from a paragraph.

    python-docx may inherit numPr through style chains even when 'Normal'
    is requested.  Explicitly deleting the element is the only safe way to
    guarantee no bullet or number character is rendered by Word.
    """
    pPr = para._p.get_or_add_pPr()
    for num_pr in pPr.findall(qn("w:numPr")):
        pPr.remove(num_pr)


def _para_bottom_border(para, color: str = "AAAAAA", sz: str = "6"):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _clean_hex(color))
    pBdr.append(bottom)
    pPr.append(pBdr)


def extract_text(tokens: List[Dict]) -> str:
    """Flatten inline tokens to plain text (used for alt-text, captions, etc.)."""
    parts = []
    for tok in tokens:
        t = tok.get("type")
        if t in ("text", "codespan"):
            parts.append(tok.get("raw", ""))
        elif t in ("softlinebreak", "linebreak"):
            parts.append(" ")
        elif "children" in tok:
            parts.append(extract_text(tok["children"]))
        elif "raw" in tok:
            parts.append(tok["raw"])
    return "".join(parts)


# ─── DocxConverter ────────────────────────────────────────────────────────────

class DocxConverter:
    def __init__(self, style: Dict, include_mermaid: bool = True):
        self.style = style
        self.include_mermaid = include_mermaid
        self.doc = Document()
        self.mermaid_images: Dict[str, bytes] = {}
        self._mermaid_idx = 0

    # ── Public entry point ────────────────────────────────────────────────────

    def convert(self, markdown_text: str) -> Tuple[bytes, Dict[str, bytes]]:
        self._setup_page()
        md = mistune.create_markdown(renderer="ast", plugins=["strikethrough", "table"])
        tokens = md(markdown_text)
        for token in tokens:
            self._block(token)
        buf = io.BytesIO()
        self.doc.save(buf)
        patched = self._patch_theme_zip(buf.getvalue())
        return patched, self.mermaid_images

    def _patch_theme_zip(self, docx_bytes: bytes) -> bytes:
        """Patch word/theme/theme1.xml inside the DOCX zip so that both
        majorEastAsia and minorEastAsia resolve to the configured font.

        python-docx's built-in template has <a:ea typeface=""/> (empty) for
        the East Asian theme font slot.  Word looks up the theme file directly
        when resolving *Theme font attributes — so even after replacing all
        w:eastAsiaTheme style references with explicit w:eastAsia names, Word
        can still fall back to the empty-typeface theme font through its own
        internal font-resolution logic, causing Korean syllable decomposition
        on the first paragraph.  Replacing the empty typeface in the theme file
        itself is the definitive fix.
        """
        fn = self.style.get("font_name", "Malgun Gothic")
        src = io.BytesIO(docx_bytes)
        dst = io.BytesIO()
        with zipfile.ZipFile(src, "r") as zin, \
             zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/theme/theme1.xml":
                    # Replace empty East Asian typeface with the configured font
                    xml = data.decode("utf-8")
                    xml = re.sub(
                        r'(<a:ea\s+typeface=")[^"]*(")',
                        rf'\g<1>{fn}\g<2>',
                        xml,
                    )
                    data = xml.encode("utf-8")
                zout.writestr(item, data)
        return dst.getvalue()

    # ── Page setup ────────────────────────────────────────────────────────────

    def _setup_page(self):
        section = self.doc.sections[0]
        if self.style.get("page_size", "A4") == "A4":
            section.page_width = Twips(11906)
            section.page_height = Twips(16838)
        else:  # Letter
            section.page_width = Twips(12240)
            section.page_height = Twips(15840)
        m = self.style.get("margin", {})
        section.top_margin    = Twips(m.get("top",    1440))
        section.bottom_margin = Twips(m.get("bottom", 1440))
        section.left_margin   = Twips(m.get("left",   1440))
        section.right_margin  = Twips(m.get("right",  1440))
        self._patch_doc_defaults()

    def _patch_doc_defaults(self):
        """Patch document-level defaults to ensure correct Korean font rendering.

        python-docx's built-in template has:
          - w:eastAsiaTheme="minorEastAsia" → resolves to empty typeface ("")
          - w:lang w:eastAsia="en-US"       → wrong language for Korean

        Word resolves East Asian font through: run rPr → paragraph mark rPr →
        character/paragraph style → docDefaults. When the docDefaults theme
        font is empty and the language is "en-US", Word may ignore per-run
        w:eastAsia settings for the very first paragraph in the document.
        Patching docDefaults with explicit values fixes this.
        """
        fn = self.style.get("font_name", "Malgun Gothic")
        styles_elem = self.doc.styles._element
        doc_defaults = styles_elem.find(qn("w:docDefaults"))
        if doc_defaults is None:
            return
        rpr_default = doc_defaults.find(qn("w:rPrDefault"))
        if rpr_default is None:
            return
        rpr = rpr_default.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            rpr_default.append(rpr)

        # Replace theme-only rFonts with explicit font names so Word never
        # falls back to the empty-typeface theme font.
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"),    fn)
        rfonts.set(qn("w:hAnsi"),    fn)
        rfonts.set(qn("w:eastAsia"), fn)

        # Fix East Asian language — "en-US" prevents proper Korean rendering.
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:eastAsia"), "ko-KR")

        # Patch all Heading* / Heading*Char built-in styles.
        # OOXML rule: when both w:eastAsia (direct) and w:eastAsiaTheme (theme)
        # are present in an effective rFonts, the THEME attribute wins.
        # The template's Heading styles use eastAsiaTheme="majorEastAsia" which
        # resolves to an empty typeface. Word applies the first outline-level-0
        # paragraph's linked Heading1Char character style, so our per-run
        # w:eastAsia gets overridden. Replacing theme refs with explicit names
        # eliminates this conflict for every Heading level.
        self._patch_all_theme_fonts(fn)

    def _patch_all_theme_fonts(self, fn: str) -> None:
        """Replace every w:eastAsiaTheme reference in styles.xml with an
        explicit w:eastAsia font name.

        The template contains many styles (Headings, Title, Subtitle, TOC
        entries, etc.) that reference majorEastAsia / minorEastAsia theme
        fonts. Both resolve to an empty typeface in python-docx's built-in
        template. Because OOXML gives theme font attributes higher priority
        than direct font attributes when both are present in effective run
        properties, these empty-typeface theme references override our
        per-run w:eastAsia="Malgun Gothic" setting and cause Korean syllable
        decomposition.  Replacing theme refs with explicit names document-wide
        eliminates the conflict at every style hierarchy level.
        """
        styles_elem = self.doc.styles._element
        theme_attrs = (
            qn("w:asciiTheme"), qn("w:hAnsiTheme"),
            qn("w:eastAsiaTheme"), qn("w:cstheme"),
        )
        for rFonts in styles_elem.iter(qn("w:rFonts")):
            # Only touch elements that carry theme font references.
            if not any(rFonts.get(a) for a in theme_attrs):
                continue
            for attr in theme_attrs:
                rFonts.attrib.pop(attr, None)
            rFonts.set(qn("w:ascii"),    fn)
            rFonts.set(qn("w:hAnsi"),    fn)
            rFonts.set(qn("w:eastAsia"), fn)

    # ── Block dispatcher ──────────────────────────────────────────────────────

    def _block(self, token: Dict, list_level: int = 0, ordered: bool = False):
        t = token.get("type")
        if   t == "heading":        self._heading(token)
        elif t == "paragraph":      self._paragraph(token)
        elif t == "block_code":     self._code_or_mermaid(token)
        elif t == "block_quote":    self._blockquote(token)
        elif t == "list":           self._list(token, level=0)
        elif t == "thematic_break": self._hr()
        elif t == "table":          self._table(token)
        elif t == "blank_line":     pass
        elif t == "block_html":     self._block_html(token)
        else:
            for child in token.get("children", []):
                self._block(child)

    # ── Headings ──────────────────────────────────────────────────────────────

    def _heading(self, token: Dict):
        level = (token.get("attrs") or {}).get("level", 1)
        sk = f"heading{min(level, 6)}"
        hs = self.style.get(sk, {})

        fn    = self.style.get("font_name", "Malgun Gothic")
        fs    = hs.get("font_size", 14)
        bold  = hs.get("bold", True)
        color = hs.get("color", "000000")

        para = self.doc.add_paragraph(style="Normal")

        # Merge all heading text into a SINGLE run.
        # mistune splits "[text]more" into two tokens: "[" + "text]more".
        # When "[" becomes a separate Latin run followed by a Korean run,
        # Word uses the Latin rendering context for the whole paragraph and
        # misrenders Korean syllables as Jamo.  A single run avoids this.
        full_text = extract_text(token.get("children", []))
        run = para.add_run(full_text)
        _set_font(run, fn)
        run.font.size  = Pt(fs)
        run.font.bold  = bold
        run.font.color.rgb = _hex(color)

        pf = para.paragraph_format
        pf.space_before = Twips(hs.get("space_before", 160))
        pf.space_after  = Twips(hs.get("space_after",   80))
        pf.keep_with_next = True
        _outline_level(para, level)

    # ── Paragraphs ────────────────────────────────────────────────────────────

    def _paragraph(self, token: Dict):
        ps = self.style.get("paragraph", {})
        para = self.doc.add_paragraph(style="Normal")
        self._inline(para, token.get("children", []))

        fn = self.style.get("font_name", "Malgun Gothic")
        fs = self.style.get("font_size_body", 11)
        for run in para.runs:
            if not run.font.name:
                _set_font(run, fn)
            if not run.font.size:
                run.font.size = Pt(fs)

        pf = para.paragraph_format
        pf.space_after  = Twips(ps.get("space_after",  120))
        pf.line_spacing = ps.get("line_spacing", 1.15)

    # ── Code blocks & mermaid ─────────────────────────────────────────────────

    def _code_or_mermaid(self, token: Dict):
        info = ((token.get("attrs") or {}).get("info") or "").strip()
        if info.lower() == "mermaid":
            self._mermaid(token.get("raw", ""))
        else:
            self._code_block(token.get("raw", ""), info)

    def _code_block(self, code: str, lang: str = ""):
        cbs = self.style.get("code_block", {})
        fn  = cbs.get("font_name", "Consolas")
        fs  = cbs.get("font_size", 9)
        bg  = cbs.get("background_color", "F2F2F2")
        fc  = cbs.get("font_color", "333333")
        sb  = cbs.get("space_before", 80)
        sa  = cbs.get("space_after", 80)

        first_para = True

        if lang:
            lp = self.doc.add_paragraph(style="Normal")
            lr = lp.add_run(lang)
            lr.font.name = fn
            lr.font.size = Pt(fs - 1)
            lr.font.italic = True
            lr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            _para_shade(lp, bg)
            lp.paragraph_format.space_before = Twips(sb)
            lp.paragraph_format.space_after  = Twips(0)
            lp.paragraph_format.left_indent  = Twips(180)
            first_para = False

        for line in code.splitlines():
            p = self.doc.add_paragraph(style="Normal")
            run = p.add_run(line or "\u00a0")  # nbsp keeps empty lines visible
            _set_font(run, fn)
            run.font.size = Pt(fs)
            run.font.color.rgb = _hex(fc)
            _para_shade(p, bg)
            p.paragraph_format.space_before = Twips(sb if first_para else 0)
            p.paragraph_format.space_after  = Twips(0)
            p.paragraph_format.left_indent  = Twips(180)
            first_para = False

        # Trailing spacer
        sp = self.doc.add_paragraph(style="Normal")
        sp.paragraph_format.space_before = Twips(0)
        sp.paragraph_format.space_after  = Twips(sa)

    # ── Blockquotes ───────────────────────────────────────────────────────────

    def _blockquote(self, token: Dict):
        fn = self.style.get("font_name", "Malgun Gothic")
        fs = self.style.get("font_size_body", 11)
        for child in token.get("children", []):
            if child.get("type") == "paragraph":
                para = self.doc.add_paragraph(style="Normal")
                self._inline(para, child.get("children", []))
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    run.font.italic = True
                    if not run.font.name:
                        _set_font(run, fn)
                    if not run.font.size:
                        run.font.size = Pt(fs)
                pf = para.paragraph_format
                pf.left_indent  = Twips(720)
                pf.space_before = Twips(60)
                pf.space_after  = Twips(60)
                _para_left_border(para, "2E75B6")
            else:
                self._block(child)

    # ── Lists ─────────────────────────────────────────────────────────────────

    def _list(self, token: Dict, level: int):
        ordered = (token.get("attrs") or {}).get("ordered", False)
        for item in token.get("children", []):
            if item.get("type") == "list_item":
                self._list_item(item, level, ordered)

    def _list_item(self, token: Dict, level: int, ordered: bool):
        show_bullet = self.style.get("show_bullet_symbol", False)
        bullet_styles  = ["List Bullet",  "List Bullet 2",  "List Bullet 3"]
        number_styles  = ["List Number",  "List Number 2",  "List Number 3"]
        style_name = (number_styles if ordered else bullet_styles)[min(level, 2)]

        fn = self.style.get("font_name", "Malgun Gothic")
        fs = self.style.get("font_size_body", 11)
        text_para_done = False

        for child in token.get("children", []):
            ct = child.get("type")
            # mistune 3.x uses "block_text" for tight lists, "paragraph" for loose lists
            if ct in ("paragraph", "block_text") and not text_para_done:
                if show_bullet:
                    try:
                        para = self.doc.add_paragraph(style=style_name)
                    except Exception:
                        para = self.doc.add_paragraph(style="Normal")
                        para.paragraph_format.left_indent = Twips(360 * (level + 1))
                else:
                    para = self.doc.add_paragraph(style="Normal")
                    para.paragraph_format.left_indent = Twips(360 * (level + 1))
                    _remove_numbering(para)
                self._inline(para, child.get("children", []))
                for run in para.runs:
                    if not run.font.name:
                        _set_font(run, fn)
                    if not run.font.size:
                        run.font.size = Pt(fs)
                text_para_done = True
            elif ct == "list":
                if not text_para_done:
                    # Empty list item — add placeholder
                    if show_bullet:
                        try:
                            para = self.doc.add_paragraph(style=style_name)
                        except Exception:
                            para = self.doc.add_paragraph(style="Normal")
                    else:
                        para = self.doc.add_paragraph(style="Normal")
                        para.paragraph_format.left_indent = Twips(360 * (level + 1))
                        _remove_numbering(para)
                    text_para_done = True
                self._list(child, level=level + 1)
            elif ct == "block_code":
                self._code_or_mermaid(child)
            else:
                self._block(child)

    # ── Horizontal rules ──────────────────────────────────────────────────────

    def _hr(self):
        para = self.doc.add_paragraph(style="Normal")
        para.paragraph_format.space_before = Twips(120)
        para.paragraph_format.space_after  = Twips(120)
        _para_bottom_border(para)

    # ── Tables ────────────────────────────────────────────────────────────────

    def _table(self, token: Dict):
        header_rows, body_rows = [], []
        for child in token.get("children", []):
            if child.get("type") == "table_head":
                # mistune 3.x: table_head children are table_cell tokens directly
                cells = child.get("children", [])
                if cells:
                    header_rows.append(cells)
            elif child.get("type") == "table_body":
                for row in child.get("children", []):
                    if row.get("type") == "table_row":
                        body_rows.append(row.get("children", []))

        all_rows = header_rows + body_rows
        if not all_rows:
            return

        n_cols = max(len(r) for r in all_rows)
        n_rows = len(all_rows)

        ths = self.style.get("table_header", {})
        alts = self.style.get("table_row_alt", {})
        h_bg   = ths.get("background_color", "2E75B6")
        h_fg   = ths.get("font_color", "FFFFFF")
        h_bold = ths.get("bold", True)
        alt_bg = alts.get("background_color", "EBF3FB")
        fn = self.style.get("font_name", "Malgun Gothic")
        fs = self.style.get("font_size_body", 11)

        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"

        for ri, row_cells in enumerate(all_rows):
            is_header = ri < len(header_rows)
            body_idx  = ri - len(header_rows)
            row = table.rows[ri]
            for ci, cell_tok in enumerate(row_cells):
                if ci >= n_cols:
                    break
                cell = row.cells[ci]
                cell.text = ""
                para = cell.paragraphs[0]
                self._inline(para, cell_tok.get("children", []))
                for run in para.runs:
                    _set_font(run, fn)
                    run.font.size = Pt(fs)
                    if is_header:
                        run.font.bold = h_bold
                        run.font.color.rgb = _hex(h_fg)
                if is_header:
                    _cell_shade(cell, h_bg)
                elif body_idx % 2 == 1:
                    _cell_shade(cell, alt_bg)

        # Spacer after table
        sp = self.doc.add_paragraph(style="Normal")
        sp.paragraph_format.space_after = Twips(120)

    # ── Mermaid ───────────────────────────────────────────────────────────────

    def _mermaid(self, code: str):
        from mermaid_renderer import detect_diagram_type, render_mermaid

        if not self.include_mermaid:
            self._code_block(code, "mermaid")
            return

        self._mermaid_idx += 1
        diagram_type = detect_diagram_type(code)
        fname = f"mermaid_{self._mermaid_idx:02d}.png"

        try:
            png = render_mermaid(code)
            self.mermaid_images[fname] = png

            img_para = self.doc.add_paragraph(style="Normal")
            img_para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
            img_run = img_para.add_run()
            img_run.add_picture(io.BytesIO(png), width=Inches(5.5))

            cap = self.doc.add_paragraph(style="Normal")
            cap.alignment = 1
            cap_run = cap.add_run(f"Figure {self._mermaid_idx}: {diagram_type}")
            cap_run.font.italic = True
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            cap.paragraph_format.space_after = Twips(120)

        except Exception as exc:
            print(f"[converter] Mermaid embed failed: {exc!r}. Falling back to code block.")
            self._code_block(code, "mermaid")

    # ── Block HTML ────────────────────────────────────────────────────────────

    def _block_html(self, token: Dict):
        raw = token.get("raw", "")
        # Strip tags and add as plain text to preserve content
        text = re.sub(r"<[^>]+>", "", raw).strip()
        if text:
            para = self.doc.add_paragraph(style="Normal")
            run = para.add_run(text)
            _set_font(run, self.style.get("font_name", "Malgun Gothic"))
            run.font.size = Pt(self.style.get("font_size_body", 11))

    # ── Inline dispatcher ─────────────────────────────────────────────────────

    def _inline(self, para, tokens: List[Dict],
                bold=False, italic=False, code=False, strike=False):
        fn   = self.style.get("font_name", "Malgun Gothic")
        fs   = self.style.get("font_size_body", 11)
        cbs  = self.style.get("code_block", {})
        cfn  = cbs.get("font_name", "Consolas")
        cfs  = cbs.get("font_size", 9)
        ics  = self.style.get("inline_code", {})
        icfn = ics.get("font_name", "Consolas")
        icfs = ics.get("font_size", 9)
        icfc = ics.get("font_color", "C7254E")
        icbg = ics.get("background_color", "F2F2F2")

        for tok in tokens:
            t = tok.get("type")

            if t == "text":
                raw = tok.get("raw", "")
                if not raw:
                    continue
                run = para.add_run(raw)
                _set_font(run, cfn if code else fn)
                run.font.size = Pt(cfs if code else fs)
                if bold:   run.font.bold   = True
                if italic: run.font.italic = True
                if strike: run.font.strike = True

            elif t == "strong":
                self._inline(para, tok.get("children", []),
                             bold=True, italic=italic, code=code, strike=strike)

            elif t == "emphasis":
                self._inline(para, tok.get("children", []),
                             bold=bold, italic=True, code=code, strike=strike)

            elif t == "strikethrough":
                self._inline(para, tok.get("children", []),
                             bold=bold, italic=italic, code=code, strike=True)

            elif t == "codespan":
                run = para.add_run(tok.get("raw", ""))
                _set_font(run, icfn)
                run.font.size = Pt(icfs)
                run.font.color.rgb = _hex(icfc)
                if bold:   run.font.bold   = True
                if italic: run.font.italic = True
                rPr = run._r.get_or_add_rPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), _clean_hex(icbg))
                rPr.append(shd)

            elif t == "link":
                url       = (tok.get("attrs") or {}).get("url", "")
                link_text = extract_text(tok.get("children", [])) or url
                if self.style.get("include_hyperlink", True):
                    _add_hyperlink(para, link_text, url)
                else:
                    run = para.add_run(link_text)
                    _set_font(run, fn)
                    run.font.size = Pt(fs)

            elif t == "image":
                attrs = tok.get("attrs") or {}
                url   = attrs.get("url", "")
                alt   = attrs.get("alt", "")
                if url and not url.startswith(("http://", "https://", "data:")):
                    img_path = Path(url)
                    if img_path.exists():
                        try:
                            run = para.add_run()
                            run.add_picture(str(img_path), width=Inches(4))
                            continue
                        except Exception:
                            pass
                # Fallback placeholder
                run = para.add_run(f"[Image: {alt or url}]")
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                _set_font(run, fn)
                run.font.size = Pt(fs)

            elif t in ("linebreak", "softlinebreak"):
                para.add_run("\n")

            elif t in ("raw_html", "html_inline"):
                text = re.sub(r"<[^>]+>", "", tok.get("raw", "")).strip()
                if text:
                    run = para.add_run(text)
                    _set_font(run, fn)
                    run.font.size = Pt(fs)

            else:
                # Unknown token: recurse into children or emit raw text
                if "children" in tok:
                    self._inline(para, tok["children"], bold, italic, code, strike)
                elif "raw" in tok:
                    run = para.add_run(tok["raw"])
                    _set_font(run, fn)
                    run.font.size = Pt(fs)


# ─── Public API ───────────────────────────────────────────────────────────────

DEFAULT_STYLE: Dict = {
    "page_size": "A4",
    "font_name": "Malgun Gothic",
    "font_size_body": 11,
    "show_bullet_symbol": True,
    "heading1": {"font_size": 20, "bold": True,  "color": "1F3864", "space_before": 240, "space_after": 120},
    "heading2": {"font_size": 16, "bold": True,  "color": "2E75B6", "space_before": 200, "space_after": 80},
    "heading3": {"font_size": 13, "bold": True,  "color": "44546A", "space_before": 160, "space_after": 60},
    "heading4": {"font_size": 12, "bold": True,  "color": "44546A", "space_before": 120, "space_after": 40},
    "heading5": {"font_size": 11, "bold": True,  "color": "595959", "space_before": 100, "space_after": 30},
    "heading6": {"font_size": 11, "bold": False, "color": "7F7F7F", "space_before": 80,  "space_after": 20},
    "paragraph": {"space_after": 120, "line_spacing": 1.15},
    "code_block":  {"font_name": "Consolas", "font_size": 9, "background_color": "F2F2F2", "font_color": "333333", "space_before": 80, "space_after": 80},
    "inline_code": {"font_name": "Consolas", "font_size": 9, "background_color": "F2F2F2", "font_color": "C7254E"},
    "table_header": {"background_color": "2E75B6", "font_color": "FFFFFF", "bold": True},
    "table_row_alt": {"background_color": "EBF3FB"},
    "margin": {"top": 1440, "bottom": 1440, "left": 1440, "right": 1440},
}


def convert_markdown_to_docx(
    markdown_text: str,
    style_config: Dict,
    include_mermaid: bool = True,
) -> Tuple[bytes, Dict[str, bytes]]:
    """
    Convert markdown text to a DOCX file.

    Returns:
        (docx_bytes, mermaid_images)
        where mermaid_images is {filename: png_bytes} for all rendered diagrams.
    """
    merged = {**DEFAULT_STYLE, **style_config}
    # Deep-merge nested dicts (one level)
    for key in ("heading1", "heading2", "heading3", "heading4", "heading5", "heading6",
                "paragraph", "code_block", "table_header", "table_row_alt", "margin"):
        if key in style_config and isinstance(style_config[key], dict):
            merged[key] = {**DEFAULT_STYLE.get(key, {}), **style_config[key]}

    conv = DocxConverter(merged, include_mermaid)
    return conv.convert(markdown_text)


def parse_markdown_structure(markdown_text: str) -> List[Dict]:
    """Return mistune AST tokens for the preview endpoint."""
    md = mistune.create_markdown(renderer="ast", plugins=["strikethrough", "table"])
    return md(markdown_text)
